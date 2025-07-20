from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import json
import csv
from io import StringIO, BytesIO
from typing import Dict, Any, List
import docx
from docx import Document

class ExportService:
    
    def export_to_pdf(self, content_data: Dict[str, Any]) -> BytesIO:
        """Export content and review data to PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("Content Generation Report", title_style))
        story.append(Spacer(1, 20))
        
        # Metadata table
        metadata = content_data.get('generated_content', {}).get('metadata', {})
        meta_data = [
            ['Topic', metadata.get('topic', 'N/A')],
            ['Content Type', metadata.get('content_type', 'N/A')],
            ['Target Audience', metadata.get('target_audience', 'N/A')],
            ['Generation Time', metadata.get('generation_timestamp', 'N/A')]
        ]
        
        meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(meta_table)
        story.append(Spacer(1, 20))
        
        # Generated Content
        story.append(Paragraph("Generated Content", styles['Heading2']))
        content_text = content_data.get('generated_content', {}).get('content', '')
        story.append(Paragraph(content_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Review Results
        story.append(Paragraph("Review Results", styles['Heading2']))
        
        final_decision = content_data.get('final_decision', {})
        decision_data = [
            ['Final Decision', final_decision.get('final_decision', 'N/A')],
            ['Final Score', f"{final_decision.get('final_score', 0):.2f}"],
            ['Timestamp', final_decision.get('timestamp', 'N/A')]
        ]
        
        decision_table = Table(decision_data, colWidths=[2*inch, 4*inch])
        decision_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(decision_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_to_word(self, content_data: Dict[str, Any]) -> BytesIO:
        """Export content to Word document"""
        doc = docx.Document()
        
        # Title
        title = doc.add_heading('Content Generation Report', 0)
        title.alignment = 1  # Center alignment
        
        # Metadata
        doc.add_heading('Content Details', level=1)
        metadata = content_data.get('generated_content', {}).get('metadata', {})
        
        p = doc.add_paragraph()
        p.add_run(f"Topic: ").bold = True
        p.add_run(metadata.get('topic', 'N/A'))
        
        p = doc.add_paragraph()
        p.add_run(f"Content Type: ").bold = True
        p.add_run(metadata.get('content_type', 'N/A'))
        
        p = doc.add_paragraph()
        p.add_run(f"Target Audience: ").bold = True
        p.add_run(metadata.get('target_audience', 'N/A'))
        
        # Generated Content
        doc.add_heading('Generated Content', level=1)
        content_text = content_data.get('generated_content', {}).get('content', '')
        doc.add_paragraph(content_text)
        
        # Review Results
        doc.add_heading('Review Results', level=1)
        final_decision = content_data.get('final_decision', {})
        
        p = doc.add_paragraph()
        p.add_run(f"Final Decision: ").bold = True
        p.add_run(final_decision.get('final_decision', 'N/A'))
        
        p = doc.add_paragraph()
        p.add_run(f"Quality Score: ").bold = True
        p.add_run(f"{final_decision.get('final_score', 0):.2f}")
        
        # Summary Points
        summary = final_decision.get('summary', [])
        if summary:
            doc.add_heading('Summary Points', level=2)
            for point in summary:
                doc.add_paragraph(point, style='List Bullet')
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_to_csv(self, content_list: List[Dict[str, Any]]) -> StringIO:
        """Export multiple content records to CSV"""
        output = StringIO()
        writer = csv.writer(output)
        
        # Headers
        headers = [
            'Topic', 'Content Type', 'Target Audience', 'Final Decision',
            'Final Score', 'Factuality Score', 'Style Score', 'Generation Time'
        ]
        writer.writerow(headers)
        
        # Data rows
        for content in content_list:
            metadata = content.get('generated_content', {}).get('metadata', {})
            final_decision = content.get('final_decision', {})
            
            row = [
                metadata.get('topic', ''),
                metadata.get('content_type', ''),
                metadata.get('target_audience', ''),
                final_decision.get('final_decision', ''),
                final_decision.get('final_score', 0),
                # Add other scores from review pipeline
                '',  # Factuality score
                '',  # Style score
                metadata.get('generation_timestamp', '')
            ]
            writer.writerow(row)
        
        output.seek(0)
        return output

# Global export service
export_service = ExportService()
